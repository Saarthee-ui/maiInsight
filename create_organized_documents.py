"""
Script to create organized Word documents from extracted RAG content.
Uses python-docx to create proper Word documents.
"""
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_document_with_content(title, content_sections):
    """Create a Word document with title and content sections."""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content sections
    for section_title, section_content in content_sections:
        if section_title:
            doc.add_heading(section_title, level=1)
        
        # Split content into paragraphs
        paragraphs = section_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.space_after = Pt(6)
        
        # Add spacing between sections
        doc.add_paragraph()
    
    return doc

def main():
    """Create organized documents from extracted content."""
    # Read extracted content
    with open("rag_documents_extracted.json", "r", encoding="utf-8") as f:
        extracted = json.load(f)
    
    # Create output directories
    Path("rag_documents/documentation").mkdir(exist_ok=True)
    Path("rag_documents/schemas").mkdir(exist_ok=True)
    Path("rag_documents/examples").mkdir(exist_ok=True)
    Path("rag_documents/rules").mkdir(exist_ok=True)
    
    print("Creating organized documents...\n")
    
    # 1. DATA WAREHOUSE OVERVIEW
    print("1. Creating data_warehouse_overview.docx...")
    overview_content = [
        ("SYSTEM ARCHITECTURE", extracted["RAG 3 SYSTEM ARCHITECTURE.docx"]["content"]),
        ("MEDALLION ARCHITECTURE", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL001]")[0] + extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL001]")[1].split("[MDL002]")[0]),
        ("DATA FLOW", "Bronze → Silver: Continuous ingestion, CDC-based raw replication. AI monitors schema and maps to hubs and satellites.\n\nSilver → Gold: Business-driven aggregation and dimensional modeling. AI determines which gold models to create based on usage and performance analytics."),
        ("GOVERNANCE & QUALITY", "Lineage captured from Bronze through Gold. Human validation at each layer. Incremental loads validated via timestamps. Schema evolution tracked in metadata vault.")
    ]
    doc = create_document_with_content("Data Warehouse Overview", overview_content)
    doc.save("rag_documents/documentation/data_warehouse_overview.docx")
    print("   ✓ Created")
    
    # 2. TRANSFORMATION GUIDE
    print("2. Creating transformation_guide.docx...")
    transform_content = [
        ("BRONZE LAYER TRANSFORMATION", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL002]")[1].split("[MDL003]")[0]),
        ("SILVER LAYER TRANSFORMATION", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL003]")[1].split("[MDL004]")[0]),
        ("GOLD LAYER TRANSFORMATION", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL004]")[1].split("[MDL005]")[0]),
        ("INCREMENTAL LOADING", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD005]")[1].split("[TD006]")[0]),
        ("DATA QUALITY CHECKS", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD007]")[1].split("[TD008]")[0])
    ]
    doc = create_document_with_content("Transformation Guide", transform_content)
    doc.save("rag_documents/documentation/transformation_guide.docx")
    print("   ✓ Created")
    
    # 3. DATABASE SCHEMAS
    print("3. Creating database_schemas.docx...")
    schema_content = [
        ("BRONZE LAYER SCHEMA", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL002]")[1].split("[MDL003]")[0]),
        ("SILVER LAYER SCHEMA", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL003]")[1].split("[MDL004]")[0]),
        ("GOLD LAYER SCHEMA", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL004]")[1].split("[MDL005]")[0])
    ]
    doc = create_document_with_content("Database Schemas", schema_content)
    doc.save("rag_documents/schemas/database_schemas.docx")
    print("   ✓ Created")
    
    # 4. TABLE DESCRIPTIONS
    print("4. Creating table_descriptions.docx...")
    table_content = [
        ("BRONZE LAYER TABLES", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL002]")[1].split("[MDL003]")[0]),
        ("SILVER LAYER TABLES", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL003]")[1].split("[MDL004]")[0]),
        ("GOLD LAYER TABLES", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL004]")[1].split("[MDL005]")[0])
    ]
    doc = create_document_with_content("Table Descriptions", table_content)
    doc.save("rag_documents/schemas/table_descriptions.docx")
    print("   ✓ Created")
    
    # 5. TRANSFORMATION EXAMPLES
    print("5. Creating transformation_examples.docx...")
    examples_content = [
        ("BRONZE TO SILVER EXAMPLE", extracted["RAG2 CONTEXTUAL MEMORY.docx"]["content"].split("[BSG006]")[1].split("[BSG007]")[0]),
        ("INCREMENTAL LOAD EXAMPLE", extracted["RAG2 CONTEXTUAL MEMORY.docx"]["content"].split("[BSG007]")[1].split("[BSG008]")[0]),
        ("SILVER TO GOLD EXAMPLE", extracted["RAG2 CONTEXTUAL MEMORY.docx"]["content"].split("[BSG004]")[1].split("[BSG005]")[0])
    ]
    doc = create_document_with_content("Transformation Examples", examples_content)
    doc.save("rag_documents/examples/transformation_examples.docx")
    print("   ✓ Created")
    
    # 6. USE CASES
    print("6. Creating use_cases.docx...")
    use_cases_content = [
        ("BUSINESS QUESTIONS FOR GOLD LAYER", extracted["RAG2b DATABASE SCHEMAS.docx"]["content"].split("[MDL005]")[1].split("[MDL006]")[0]),
        ("AI USE-CASES IN TELECOM", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD004]")[1].split("[TD005]")[0]),
        ("CUSTOMER SEGMENTATION", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD006]")[1].split("[TD007]")[0])
    ]
    doc = create_document_with_content("Use Cases", use_cases_content)
    doc.save("rag_documents/examples/use_cases.docx")
    print("   ✓ Created")
    
    # 7. NAMING CONVENTIONS
    print("7. Creating naming_conventions.docx...")
    naming_content = [
        ("BRONZE LAYER NAMING", "Use suffix _BZ for all Bronze layer tables.\nExample: CUSTOMER_BZ, ORDER_BZ, PRODUCT_BZ"),
        ("SILVER LAYER NAMING", "Use prefixes HUB_, SAT_, LINK_ for Data Vault structures.\nExample: HUB_CUSTOMER, SAT_CUSTOMER_ATTR, LINK_CUSTOMER_ORDER"),
        ("GOLD LAYER NAMING", "Use prefixes FACT_, DIM_ for dimensional model.\nExample: FACT_ORDER, DIM_CUSTOMER, DIM_TIME"),
        ("TRANSFORMATION NAMING", extracted["PROCESS.docx"]["content"].split("transformation name")[1].split("Suggest")[0] if "transformation name" in extracted["PROCESS.docx"]["content"] else "Use format: [PURPOSE]_[DATABASE]_[TYPE]\nExample: SALES_DASHBOARD_DAILY, CUSTOMER_ANALYTICS_MONTHLY")
    ]
    doc = create_document_with_content("Naming Conventions", naming_content)
    doc.save("rag_documents/rules/naming_conventions.docx")
    print("   ✓ Created")
    
    # 8. BUSINESS RULES
    print("8. Creating business_rules.docx...")
    business_rules_content = [
        ("BRONZE-SILVER-GOLD ARCHITECTURE RULES", extracted["RAG2 CONTEXTUAL MEMORY.docx"]["content"].split("[BSG001]")[1].split("[BSG002]")[0]),
        ("ADAPTIVE GOLD LAYER STRATEGY", extracted["RAG2 CONTEXTUAL MEMORY.docx"]["content"].split("[BSG002]")[1].split("[BSG003]")[0]),
        ("GOVERNANCE AND COMPLIANCE", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD008]")[1].split("[TD009]")[0])
    ]
    doc = create_document_with_content("Business Rules", business_rules_content)
    doc.save("rag_documents/rules/business_rules.docx")
    print("   ✓ Created")
    
    # 9. DATA QUALITY RULES
    print("9. Creating data_quality_rules.docx...")
    quality_rules_content = [
        ("DATA QUALITY CHALLENGES", extracted["RAG 1 DOMAIN KNOWLEDGE.docx"]["content"].split("[TD007]")[1].split("[TD008]")[0]),
        ("DATA QUALITY SOLUTIONS", "Solutions: data quality rules, schema validation, anomaly detection, correction workflows, and lineage visualization dashboards."),
        ("REQUIRED FIELDS", "CUSTOMER_ID: Must not be NULL\nORDER_DATE: Must be valid date\nTOTAL_AMOUNT: Must be >= 0"),
        ("VALIDATION RULES", "Email format: Must match pattern\nPhone numbers: Must be 10 digits\nDates: Must be between valid range")
    ]
    doc = create_document_with_content("Data Quality Rules", quality_rules_content)
    doc.save("rag_documents/rules/data_quality_rules.docx")
    print("   ✓ Created")
    
    # 10. BUILD PROCESS GUIDE
    print("10. Creating build_process_guide.docx...")
    build_process_content = [
        ("BUILD STAGE OVERVIEW", extracted["PROCESS.docx"]["content"].split("Build Stage")[1].split("buildRetrieval Agent")[0] if "Build Stage" in extracted["PROCESS.docx"]["content"] else ""),
        ("BUILD RETRIEVAL AGENT", extracted["PROCESS.docx"]["content"].split("buildRetrieval Agent")[1].split("buildActionFolder")[0] if "buildRetrieval Agent" in extracted["PROCESS.docx"]["content"] else ""),
        ("ANALYTICS PROCESS", extracted["PROCESS.docx"]["content"].split("Analytics Process")[1] if "Analytics Process" in extracted["PROCESS.docx"]["content"] else "")
    ]
    doc = create_document_with_content("Build Process Guide", build_process_content)
    doc.save("rag_documents/documentation/build_process_guide.docx")
    print("   ✓ Created")
    
    print("\n✅ All documents created successfully!")
    print("\nCreated files:")
    print("  documentation/")
    print("    - data_warehouse_overview.docx")
    print("    - transformation_guide.docx")
    print("    - build_process_guide.docx")
    print("  schemas/")
    print("    - database_schemas.docx")
    print("    - table_descriptions.docx")
    print("  examples/")
    print("    - transformation_examples.docx")
    print("    - use_cases.docx")
    print("  rules/")
    print("    - naming_conventions.docx")
    print("    - business_rules.docx")
    print("    - data_quality_rules.docx")

if __name__ == "__main__":
    main()

